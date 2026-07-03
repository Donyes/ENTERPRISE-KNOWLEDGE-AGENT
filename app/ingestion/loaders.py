from pathlib import Path
from typing import List

from langchain_core.documents import Document
from pypdf import PdfReader
from docx import Document as WordDocument


SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx"}


def load_txt_or_md(file_path: Path) -> List[Document]:
    """
    Load a txt or markdown file as one LangChain Document.
    """
    text = file_path.read_text(encoding="utf-8")

    return [
        Document(
            page_content=text,
            metadata={
                "source": str(file_path),
                "file_name": file_path.name,
                "file_type": file_path.suffix.lower(),
            },
        )
    ]


def load_pdf(file_path: Path) -> List[Document]:
    """
    Load a PDF file page by page.

    Each PDF page becomes one LangChain Document.
    """
    reader = PdfReader(str(file_path))
    documents = []

    for page_index, page in enumerate(reader.pages):
        text = page.extract_text() or ""

        if not text.strip():
            continue

        documents.append(
            Document(
                page_content=text,
                metadata={
                    "source": str(file_path),
                    "file_name": file_path.name,
                    "file_type": file_path.suffix.lower(),
                    "page_number": page_index + 1,
                },
            )
        )

    return documents


def load_docx(file_path: Path) -> List[Document]:
    """
    Load a Word document as one LangChain Document.
    """
    word_doc = WordDocument(str(file_path))

    paragraphs = []
    for paragraph in word_doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            paragraphs.append(text)

    full_text = "\n".join(paragraphs)

    return [
        Document(
            page_content=full_text,
            metadata={
                "source": str(file_path),
                "file_name": file_path.name,
                "file_type": file_path.suffix.lower(),
            },
        )
    ]


def load_single_document(file_path: Path) -> List[Document]:
    """
    Load one supported file and return LangChain Document objects.
    """
    suffix = file_path.suffix.lower()

    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file type: {suffix}")

    if suffix in {".txt", ".md"}:
        return load_txt_or_md(file_path)

    if suffix == ".pdf":
        return load_pdf(file_path)

    if suffix == ".docx":
        return load_docx(file_path)

    raise ValueError(f"Unsupported file type: {suffix}")


def load_documents_from_directory(directory: str) -> List[Document]:
    """
    Load all supported documents from a directory.
    """
    directory_path = Path(directory)

    if not directory_path.exists():
        raise FileNotFoundError(f"Directory not found: {directory}")

    all_documents = []

    for file_path in directory_path.rglob("*"):
        if file_path.is_file() and file_path.suffix.lower() in SUPPORTED_EXTENSIONS:
            documents = load_single_document(file_path)
            all_documents.extend(documents)

    return all_documents